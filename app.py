import streamlit as st
from google import genai 
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import os
import re 
import requests 
from PIL import Image
import datetime
import matplotlib
import pandas as pd

# --- PREVENT GRAPHING FREEZE ---
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# --- 0. VERSION CONTROL ---
APP_VERSION = "v3.10.16-StableRollback"
DB_FILE = "zyn_database.json"

# --- 1. CONFIGURATION ---
st.set_page_config(page_title=f"ZyntaxAI {APP_VERSION}", layout="wide", page_icon="🏗️")

# Pulling keys from the hidden cloud safe!
API_KEY = st.secrets["GEMINI_API_KEY"]
MAPS_API_KEY = st.secrets["MAPS_API_KEY"]

# INITIALIZE NEW GOOGLE CLIENT
try:
    client = genai.Client(api_key=API_KEY)
except Exception as e:
    st.error(f"API Key Error: {e}")

# --- 2. DATABASE FUNCTIONS ---
def load_db():
    if not os.path.exists(DB_FILE): return {}
    try:
        with open(DB_FILE, 'r') as f: return json.load(f)
    except: return {}

def save_to_db(address, data):
    db = load_db()
    key = address.strip().lower()
    data['last_updated'] = str(datetime.date.today())
    db[key] = data
    with open(DB_FILE, 'w') as f: json.dump(db, f, indent=4)

def get_from_db(address):
    db = load_db()
    return db.get(address.strip().lower())

# --- 3. AUTO-DETECT MODELS ---
def get_models():
    try:
        model_names = []
        for m in client.models.list():
            name = m.name
            if "gemini" in name and ("flash" in name or "pro" in name):
                if "tts" not in name and "exp" not in name and "8b" not in name and "thinking" not in name:
                    model_names.append(name)
        if not model_names: 
            return ["gemini-2.5-flash", "gemini-1.5-flash", "gemini-1.5-pro"]
        return sorted(list(set(model_names)), reverse=True)
    except: 
        return ["gemini-2.5-flash", "gemini-1.5-flash", "gemini-1.5-pro"]

available_models = get_models()
idx = 0
for i, m in enumerate(available_models):
    if m == "gemini-2.5-flash" or m == "models/gemini-2.5-flash": 
        idx = i; break
    elif m == "gemini-1.5-flash" or m == "models/gemini-1.5-flash":
        idx = i

st.sidebar.header("⚙️ System Settings")
selected_model_name = st.sidebar.selectbox("Select AI Model:", available_models, index=idx)

if 'analysis_data' not in st.session_state: st.session_state['analysis_data'] = None
if 'address_cache' not in st.session_state: st.session_state['address_cache'] = ""

# --- 4. GOOGLE MAPS AUTOMATION ---
def get_google_maps_images(address):
    """Fetches Satellite and Street View images automatically from Google."""
    try:
        sat_url = f"https://maps.googleapis.com/maps/api/staticmap?center={address}&zoom=20&size=640x640&maptype=satellite&key={MAPS_API_KEY}"
        sat_response = requests.get(sat_url)
        sat_img = Image.open(io.BytesIO(sat_response.content))

        street_url = f"https://maps.googleapis.com/maps/api/streetview?size=640x640&location={address}&key={MAPS_API_KEY}"
        street_response = requests.get(street_url)
        street_img = Image.open(io.BytesIO(street_response.content))

        return sat_img, street_img
    except Exception as e:
        st.error(f"Error fetching Google Images: {e}")
        return None, None

# --- 5. THE PRO LOGIC ---
SYSTEM_PROMPT = """
You are ZyntaxAI, a Senior Property Analyst.
Analyze the provided images (Satellite/Street View) and Address.

YOUR MISSION:
1. Analyze the Context provided by the user.
2. Scan the Satellite image for risks.
3. Calculate strict Build Costs.
4. Estimate the Maximum Legally Allowable Units.
5. Calculate the Maximum Walk-Away Price (Residual Land Value).

Output STRICT JSON. DO NOT use markdown.

{
  "confidence_score": "95%",
  "confidence_explanation": "Reason for score.",
  "address_details": { "full_address": "Str" },
  "user_context_summary": "Brief summary of user notes.",
  
  "site_fundamentals": [
      {"field": "Council", "val": "Text", "imp": "Text"},
      {"field": "Zoning", "val": "Text", "imp": "Text"},
      {"field": "Dimensions", "val": "Text", "imp": "Text"},
      {"field": "Topography", "val": "Text", "imp": "Text"},
      {"field": "Premium Drivers", "val": "Text", "imp": "Text"},
      {"field": "Est. Value", "val": "$X", "imp": "Comment on value"}
  ],

  "automated_risk_scan": {
      "bushfire_risk": "Low/Medium/High (Evidence)",
      "flood_risk": "Low/Medium/High (Evidence)",
      "slope_risk": "Low/Medium/High (Evidence)",
      "easement_risk": "Low/Medium/High (Likelihood)"
  },

  "construction_cost_assumptions": "Explicitly state: 'Calculated at $X/sqm based on [State] rates...'",

  "residual_land_value": {
      "gross_realisation_value": "$Amount",
      "target_margin": "$Amount",
      "estimated_build_area_sqm": "Number (Total internal area for all units combined, e.g. 240)",
      "build_cost_per_sqm": "$Amount",
      "construction_and_soft_costs": "$Amount",
      "max_walk_away_price": "$Amount",
      "verdict": "COMPARE the Max Walk-Away Price to the Est. Value. If Walk-Away is lower, declare it UNFEASIBLE."
  },

  "financial_matrix": [
      {"strategy": "Strategy A: Hold (Rent)", "cash": "$Amount", "sqm_cost": "N/A", "yield": "X.X%", "profit": "$Amount", "risk": "Low"},
      {"strategy": "Strategy B: Renovate/Flip", "cash": "$Amount", "sqm_cost": "$Amount", "yield": "X.X%", "profit": "$Amount", "risk": "Medium"},
      {"strategy": "Strategy C: Duplex (1-into-2)", "cash": "$Amount", "sqm_cost": "$Amount", "yield": "15.0%", "profit": "$Amount", "risk": "High"},
      {"strategy": "Strategy D: Townhouses (X Units)", "cash": "$Amount", "sqm_cost": "$Amount", "yield": "20.0%", "profit": "$Amount", "risk": "High"}
  ],
  "matrix_explanation": "Concise summary.",

  "wealth_gap_data": {
      "years": [0, 1, 2, 3, 4, 5],
      "strategy_a_values": [1.4, 1.45, 1.5, 1.55, 1.6, 1.65],
      "strategy_dev_name": "Strategy D: Townhouses (X Units)",
      "strategy_dev_values": [1.4, 1.5, 2.5, 3.2, 3.3, 3.5],
      "explanation": "Explain graph."
  },

  "schematic_data": {
      "lot_width_m": 15.24, "lot_depth_m": 48.0, "max_yield_units": 3, "explanation": "Explain layout options."
  },

  "narrative": {
      "verdict_status": "HIGH RISK / MEDIUM RISK / LOW RISK", "verdict_summary": "Professional summary.", "constraints": ["Constraint 1"]
  }
}

RULES:
1. Base Range by State: NSW(2400-4300), ACT(2400-4200), VIC(2200-3900), QLD(2100-3800), TAS(2200-4200), NT(2300-4500), WA(1900-3600), SA(1700-3200).
2. Base rate = Bottom of range (Standard), Middle (Medium), Top (Luxury). Add $300/sqm if sloping. Assign this to "build_cost_per_sqm".
3. "max_yield_units" represents the absolute maximum townhouses that can fit. If max is 2, OMIT Strategy D.
4. GRV REALITY CHECK: Brand new units sell for a premium. Forecast an aggressive GRV for the completed units.
5. FEASIBILITY REALITY CHECK: Work backward from GRV. If Walk-Away Price is lower than Est. Value, your verdict MUST state: "UNFEASIBLE: The maximum allowable land price ($X) is lower than the current market value ($Y)."
"""

UPDATE_PROMPT = """
You are ZyntaxAI. You previously generated this property analysis:
{existing_json}
The user has UPDATED their Context/Notes to: "{new_context}"
Target Build Finish: "{build_finish}"
YOUR MISSION:
Recalculate the "construction_cost_assumptions", "residual_land_value", "financial_matrix", and "wealth_gap_data" based ONLY on the new context and build finish.
CRITICAL: You MUST keep "site_fundamentals", "automated_risk_scan", and "schematic_data" EXACTLY as they are. 
Output STRICT JSON.
"""

# --- 6. GENERATORS (CHARTS, DOC, AND OPEN-BOOK EXCEL) ---
def extract_number(text_val):
    try:
        clean = re.sub(r'[^\d.]', '', str(text_val))
        return float(clean) if clean else 0.0
    except:
        return 0.0

def create_excel_report(data):
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        
        format_title = workbook.add_format({'bold': True, 'font_size': 14, 'color': '#2C3E50'})
        format_help = workbook.add_format({'italic': True, 'color': '#155724', 'bg_color': '#D4EDDA', 'border': 1, 'text_wrap': True})
        format_input = workbook.add_format({'bg_color': '#FFF2CC', 'border': 1, 'num_format': '$#,##0'})
        format_input_sqm = workbook.add_format({'bg_color': '#FFF2CC', 'border': 1, 'num_format': '#,##0'})
        format_formula = workbook.add_format({'bg_color': '#D9E1F2', 'border': 1, 'bold': True, 'num_format': '$#,##0'})
        format_label = workbook.add_format({'bold': True, 'border': 1, 'bg_color': '#E7E6E6'})
        format_percent = workbook.add_format({'bg_color': '#FFF2CC', 'border': 1, 'num_format': '0.0%'})
        format_static = workbook.add_format({'border': 1})

        base_land_cost = 0.0
        for item in data.get('site_fundamentals', []):
            if item.get('field', '').lower() == 'est. value':
                base_land_cost = extract_number(item.get('val', '0'))
                break
        if base_land_cost == 0:
            base_land_cost = extract_number(data.get('residual_land_value', {}).get('max_walk_away_price', 0))

        worksheet1 = workbook.add_worksheet('Walk-Away Calculator')
        worksheet1.set_column('A:A', 40); worksheet1.set_column('B:B', 20)
        
        worksheet1.write('A1', 'ZYNTAX AI: RESIDUAL LAND VALUE CALCULATOR', format_title)
        worksheet1.merge_range('A2:B3', 'HOW TO USE: Cells highlighted in YELLOW are your assumptions. You can change these numbers to match your quotes. Cells in BLUE are formulas and will recalculate automatically.', format_help)
        
        rlv = data.get('residual_land_value', {})
        grv = extract_number(rlv.get('gross_realisation_value', 0))
        area_sqm = extract_number(rlv.get('estimated_build_area_sqm', 200))
        cost_sqm = extract_number(rlv.get('build_cost_per_sqm', 2000))
        
        worksheet1.write('A5', '1. Gross Realisation Value (GRV)', format_label)
        worksheet1.write_number('B5', grv, format_input)
        worksheet1.write('A6', '2. Target Developer Margin (%)', format_label)
        worksheet1.write_number('B6', 0.20, format_percent)
        worksheet1.write('A7', '=> Target Profit ($)', format_label)
        worksheet1.write_formula('B7', '=B5*B6', format_formula)
        worksheet1.write('A9', '3. Estimated Total Build Area (sqm)', format_label)
        worksheet1.write_number('B9', area_sqm, format_input_sqm)
        worksheet1.write('A10', '4. Base Build Cost per SQM ($)', format_label)
        worksheet1.write_number('B10', cost_sqm, format_input)
        worksheet1.write('A11', '5. Turnkey & Soft Costs Buffer (%)', format_label)
        worksheet1.write_number('B11', 0.15, format_percent)
        worksheet1.write('A12', '=> Total Dev & Soft Costs', format_label)
        worksheet1.write_formula('B12', '=B9*B10*(1+B11)', format_formula)
        worksheet1.write('A14', 'MAXIMUM WALK-AWAY LAND PRICE', format_label)
        worksheet1.write_formula('B14', '=B5-B7-B12', format_formula)

        worksheet2 = workbook.add_worksheet('Decision Matrix')
        worksheet2.set_column('A:A', 30); worksheet2.set_column('B:F', 20); worksheet2.set_column('G:G', 15)
        
        worksheet2.write('A1', 'STRATEGY COMPARISON MATRIX (OPEN BOOK)', format_title)
        worksheet2.merge_range('A2:G3', 'HOW TO USE: Yellow cells are your variables. Change the Build Area or End Sale Value to test your own scenarios. The Build Cost/sqm is live-linked to Tab 1.', format_help)
        
        worksheet2.write('A5', 'Assumed Land Purchase Price ($):', format_label)
        worksheet2.write_number('B5', base_land_cost, format_input)
        
        headers = ['Strategy', 'End Sale Value (GRV)', 'Build Area (sqm)', 'Build Cost/sqm', 'Total Cash Outlay', 'Net Profit', 'Risk']
        for col_num, header in enumerate(headers):
            worksheet2.write(6, col_num, header, format_label)
            
        row_num = 7
        for row_data in data.get('financial_matrix', []):
            strat = row_data.get('strategy', '')
            cash_val = extract_number(row_data.get('cash', 0))
            profit_val = extract_number(row_data.get('profit', 0))
            sqm_cost_val = extract_number(row_data.get('sqm_cost', 0))

            worksheet2.write(row_num, 0, strat, format_static)

            if sqm_cost_val > 500 and base_land_cost > 0 and cash_val > base_land_cost:
                build_cash = cash_val - base_land_cost
                implied_area = build_cash / (sqm_cost_val * 1.15)
                original_grv = profit_val + cash_val

                worksheet2.write_number(row_num, 1, original_grv, format_input) 
                worksheet2.write_number(row_num, 2, implied_area, format_input_sqm) 
                worksheet2.write_formula(row_num, 3, "='Walk-Away Calculator'!$B$10", format_formula) 
                
                worksheet2.write_formula(row_num, 4, f"=$B$5+(C{row_num+1}*D{row_num+1}*(1+'Walk-Away Calculator'!$B$11))", format_formula) 
                worksheet2.write_formula(row_num, 5, f"=B{row_num+1}-E{row_num+1}", format_formula) 
                worksheet2.write_string(row_num, 6, row_data.get('risk', ''), format_static)
            else:
                worksheet2.write_string(row_num, 1, "-", format_static)
                worksheet2.write_string(row_num, 2, "-", format_static)
                worksheet2.write_string(row_num, 3, "-", format_static)
                worksheet2.write_number(row_num, 4, cash_val, format_input)
                worksheet2.write_number(row_num, 5, profit_val, format_input)
                worksheet2.write_string(row_num, 6, row_data.get('risk', ''), format_static)
                
            row_num += 1

    return output.getvalue()

def generate_wealth_chart(data):
    years = data.get('years', [0,1,2,3,4,5])
    strat_a = data.get('strategy_a_values', [1,1,1,1,1,1])
    strat_dev = data.get('strategy_dev_values', [1,1,1,1,1,1])
    dev_name = data.get('strategy_dev_name', 'Strategy: Develop')

    plt.figure(figsize=(6, 4))
    plt.plot(years, strat_a, label='Strategy A: Hold', linestyle='--', color='gray', marker='o')
    plt.plot(years, strat_dev, label=dev_name, linewidth=2, color='green', marker='o')
    plt.title('5-Year Wealth Gap'); plt.ylabel('Asset Value ($M)'); plt.xlabel('Years')
    plt.grid(True, linestyle='--', alpha=0.5); plt.legend(); plt.tight_layout()
    img_buf = io.BytesIO(); plt.savefig(img_buf, format='png', dpi=100); plt.close()
    return img_buf

def generate_schematic(data, units=2):
    width = float(data.get('lot_width_m', 15.24)); depth = float(data.get('lot_depth_m', 40.0))
    fig, ax = plt.subplots(figsize=(3, 6))
    lot = plt.Rectangle((0, 0), width, depth, fill=False, edgecolor='black', linewidth=2)
    ax.add_patch(lot)
    margin_front = 6.0; margin_rear = 4.0; margin_side = 1.0
    
    if units == 2 and width >= 14.0:
        ax.plot([width/2, width/2], [0, depth], color='orange', linestyle='--', linewidth=2)
        b_width = (width/2) - (margin_side * 2); b_depth = depth * 0.45
        rect1 = plt.Rectangle((margin_side, margin_front), b_width, b_depth, facecolor='none', edgecolor='green', hatch='//')
        rect2 = plt.Rectangle((width/2 + margin_side, margin_front), b_width, b_depth, facecolor='none', edgecolor='green', hatch='//')
        ax.add_patch(rect1); ax.add_patch(rect2)
    else:
        driveway_w = 3.0; b_width = width - driveway_w - (margin_side * 2)
        ax.add_patch(plt.Rectangle((width - driveway_w, 0), driveway_w, depth, facecolor='lightgray', edgecolor='none', alpha=0.5))
        available_depth = depth - margin_front - margin_rear; spacing = 3.0
        b_depth = (available_depth - (spacing * (units - 1))) / units
        for i in range(units):
            y_pos = margin_front + i * (b_depth + spacing)
            rect = plt.Rectangle((margin_side, y_pos), b_width, b_depth, facecolor='none', edgecolor='green', hatch='//')
            ax.add_patch(rect)
    
    ax.set_xlim(-2, width+2); ax.set_ylim(-2, depth+2); ax.set_aspect('equal'); ax.axis('off')
    title_text = "Duplex\n(1-into-2)" if units == 2 else f"Townhouses\n(1-into-{units})"
    plt.title(f"{title_text}\n({width}m Frontage)", fontsize=10); plt.tight_layout()
    img_buf = io.BytesIO(); plt.savefig(img_buf, format='png', dpi=100); plt.close()
    return img_buf

def create_pro_report(data, address):
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Aptos'; style.font.size = Pt(10)
    
    section = doc.sections[0]; header = section.header; paragraph = header.paragraphs[0]; paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
    try: run = paragraph.add_run(); run.add_picture("logo.png", width=Inches(1.2))
    except: paragraph.text = "ZYNTAX AI"

    doc.add_heading(f"Property Risk Audit", 0); doc.add_paragraph(f"Address: {address}")
    
    conf_p = doc.add_paragraph()
    conf_run = conf_p.add_run(f"Confidence: {data.get('confidence_score', 'N/A')}"); conf_run.bold = True; conf_run.font.color.rgb = RGBColor(0, 0, 255)
    conf_p.add_run(f" | Date: {datetime.date.today()}")

    if data.get('user_context_summary'): doc.add_paragraph(f"Context Note: {data.get('user_context_summary')}", style='Intense Quote')

    doc.add_heading("1. Site Fundamentals", level=1)
    table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
    hdr = table.rows[0].cells; hdr[0].text = "Field"; hdr[1].text = "Attribute"; hdr[2].text = "Implications"
    for item in data.get('site_fundamentals', []):
        row = table.add_row().cells; row[0].text = str(item.get('field', '')); row[1].text = str(item.get('val', '')); row[2].text = str(item.get('imp', ''))
    
    doc.add_heading("2. Automated Risk Scan", level=1)
    risk_table = doc.add_table(rows=1, cols=2); risk_table.style = 'Table Grid'
    rhdr = risk_table.rows[0].cells; rhdr[0].text = "Risk Factor"; rhdr[1].text = "Assessment & Evidence"
    for key, val in data.get('automated_risk_scan', {}).items():
        row = risk_table.add_row().cells; row[0].text = key.replace('_', ' ').title(); row[1].text = str(val)
        if "High" in str(val): row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0); row[1].paragraphs[0].runs[0].bold = True

    doc.add_heading("3. Maximum Walk-Away Price (Residual Land Value)", level=1)
    rlv_data = data.get('residual_land_value', {})
    rlv_table = doc.add_table(rows=1, cols=2); rlv_table.style = 'Table Grid'
    r1 = rlv_table.rows[0].cells; r1[0].text = "Gross Realisation Value (GRV)"; r1[1].text = str(rlv_data.get('gross_realisation_value', ''))
    r2 = rlv_table.add_row().cells; r2[0].text = "Less: Target Margin (20%)"; r2[1].text = str(rlv_data.get('target_margin', ''))
    r3 = rlv_table.add_row().cells; r3[0].text = "Less: Dev & Soft Costs"; r3[1].text = str(rlv_data.get('construction_and_soft_costs', ''))
    r4 = rlv_table.add_row().cells; r4[0].text = "Maximum Land Purchase Price"; r4[1].text = str(rlv_data.get('max_walk_away_price', ''))
    r4[0].paragraphs[0].runs[0].bold = True; r4[1].paragraphs[0].runs[0].bold = True; r4[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 150, 0)
    
    verdict_text = rlv_data.get('verdict', '')
    p_verdict = doc.add_paragraph()
    v_run = p_verdict.add_run(verdict_text)
    if "UNFEASIBLE" in verdict_text.upper():
        v_run.font.color.rgb = RGBColor(255, 0, 0)
        v_run.bold = True
    p_verdict.style = 'Intense Quote'

    doc.add_heading("4. Investor Decision Matrix", level=1)
    if data.get('construction_cost_assumptions'):
        p = doc.add_paragraph(); p.add_run("Cost Assumptions: ").bold = True; p.add_run(data.get('construction_cost_assumptions')); p.style = 'Intense Quote'
        
    table2 = doc.add_table(rows=1, cols=6); table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Strategy"; hdr2[1].text = "Total Cash"; hdr2[2].text = "$/sqm"; hdr2[3].text = "Yield"; hdr2[4].text = "Profit"; hdr2[5].text = "Risk"
    for row_data in data.get('financial_matrix', []):
        row = table2.add_row().cells
        row[0].text = str(row_data.get('strategy', '')); row[1].text = str(row_data.get('cash', '')); row[2].text = str(row_data.get('sqm_cost', '')); 
        row[3].text = str(row_data.get('yield', '')); row[4].text = str(row_data.get('profit', '')); row[5].text = str(row_data.get('risk', ''))
    doc.add_paragraph("\n"); doc.add_paragraph(data.get('matrix_explanation', ''))

    doc.add_heading("5. Visual Analytics", level=1)
    wealth_chart = generate_wealth_chart(data.get('wealth_gap_data', {}))
    doc.add_picture(wealth_chart, width=Inches(5.5)); doc.add_paragraph("Figure 1: 5-Year Wealth Gap", style='Caption')
    
    max_units = int(data.get('schematic_data', {}).get('max_yield_units', 2))
    vis_table = doc.add_table(rows=1, cols=2 if max_units > 2 else 1)
    cell1 = vis_table.rows[0].cells[0]; p1 = cell1.add_paragraph()
    p1.add_run().add_picture(generate_schematic(data.get('schematic_data', {}), 2), width=Inches(2.5))
    cell1.add_paragraph("Layout A: Duplex Option", style='Caption')
    
    if max_units > 2:
        cell2 = vis_table.rows[0].cells[1]; p2 = cell2.add_paragraph()
        p2.add_run().add_picture(generate_schematic(data.get('schematic_data', {}), max_units), width=Inches(2.5))
        cell2.add_paragraph(f"Layout B: Townhouses ({max_units} Units)", style='Caption')

    doc.add_paragraph("\n"); doc.add_paragraph(data.get('schematic_data', {}).get('explanation', ''))

    doc.add_heading("6. Analyst Verdict", level=1)
    verdict = data.get('narrative', {}); p = doc.add_paragraph()
    run = p.add_run(f"VERDICT: {data.get('narrative', {}).get('verdict_status', 'UNKNOWN')}")
    run.bold = True; run.font.color.rgb = RGBColor(255, 0, 0) if "HIGH" in data.get('narrative', {}).get('verdict_status', '') else RGBColor(0, 150, 0)
    doc.add_paragraph(verdict.get('verdict_summary', ''))
    
    doc.add_heading("Critical Constraints:", level=2)
    for c in verdict.get('constraints', []): doc.add_paragraph(c, style='List Bullet')

    disclaimer = doc.add_paragraph("\nDisclaimer: Preliminary audit. Does not constitute financial advice. Confirm with professionals.")
    for run in disclaimer.runs: run.font.italic = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(100, 100, 100)
    
    target_stream = io.BytesIO(); doc.save(target_stream); return target_stream.getvalue()

# --- UI ---
st.title(f"ZyntaxAI | {APP_VERSION}")

col1, col2 = st.columns([1, 2])
with col1:
    address_input = st.text_input("Target Address", placeholder="e.g. 12 Smith St")
    context_input = st.text_area("Project Context / Notes", height=80, placeholder="e.g. Price guide $850k. Zoned R3.")
    build_finish = st.selectbox("Target Build Finish", ["Standard (Investment Grade)", "Medium (Custom Spec)", "Luxury (Owner Occupier)"])
    
    st.markdown("### Images Automated 📡")
    st.info("Satellite and Street View images will be fetched automatically when you run the scan.")
    
    st.divider()
    is_in_db = False
    if address_input:
        existing = get_from_db(address_input)
        if existing:
            is_in_db = True
            st.success("✅ Property in Database")
    
    if is_in_db:
        c_btn1, c_btn2, c_btn3 = st.columns(3)
        with c_btn1: run_btn = st.button("📂 LOAD", type="primary", use_container_width=True)
        with c_btn2: recalc_btn = st.button("⚡ RECALC", type="primary", use_container_width=True)
        with c_btn3: force_btn = st.button("🔄 FULL RUN", type="secondary", use_container_width=True)
    else:
        c_btn1, c_btn2 = st.columns(2)
        with c_btn1: run_btn = st.button("🚀 RUN AI ANALYSIS", type="primary", use_container_width=True)
        with c_btn2: force_btn = st.button("🔄 FORCE RUN", type="secondary", use_container_width=True)
        recalc_btn = False

    if st.button("Reset"): st.session_state['analysis_data'] = None; st.session_state['address_cache'] = ""; st.rerun()

with col2:
    if run_btn and address_input and is_in_db:
        data = get_from_db(address_input)
        st.session_state['analysis_data'] = data; st.session_state['address_cache'] = address_input; st.toast("Loaded from DB", icon="📂")

    elif recalc_btn and address_input and is_in_db:
        with st.spinner(f"Fast Recalculating Financials..."):
            try:
                existing_data = get_from_db(address_input)
                formatted_prompt = UPDATE_PROMPT.format(
                    existing_json=json.dumps(existing_data),
                    new_context=context_input,
                    build_finish=build_finish
                )
                response = client.models.generate_content(
                    model=selected_model_name,
                    contents=[formatted_prompt]
                )
                txt = response.text
                j_start = txt.find('{'); j_end = txt.rfind('}') + 1
                new_data = json.loads(txt[j_start:j_end])
                
                save_to_db(address_input, new_data)
                st.session_state['analysis_data'] = new_data; st.session_state['address_cache'] = address_input; st.toast("Financials Updated", icon="⚡")
            except Exception as e: st.error(f"Recalc Error: {e}")

    elif (run_btn and not is_in_db) or force_btn:
        if address_input:
            with st.spinner(f"Fetching Maps & Running Risk Scan..."):
                try:
                    sat_img, street_img = get_google_maps_images(address_input)
                    if sat_img and street_img:
                        st.image([sat_img, street_img], width=300, caption=["Satellite View", "Street View"])
                        
                        img_payload = [sat_img, street_img]
                        text_prompt = f"{SYSTEM_PROMPT}\n\nAddress: {address_input}.\nUser Context Note: {context_input}\nTarget Build Finish: {build_finish}"
                        contents = [text_prompt] + img_payload
                        
                        response = client.models.generate_content(
                            model=selected_model_name,
                            contents=contents
                        )
                        txt = response.text
                        j_start = txt.find('{'); j_end = txt.rfind('}') + 1
                        new_data = json.loads(txt[j_start:j_end])
                        
                        save_to_db(address_input, new_data)
                        st.session_state['analysis_data'] = new_data; st.session_state['address_cache'] = address_input; st.toast("Saved to DB", icon="💾")
                    else:
                        st.error("Failed to fetch images from Google Maps. Check your Address or Maps API Key.")
                except Exception as e: st.error(f"Error: {e}")
        else: st.warning("Please enter a Target Address to run the scan.")

    if st.session_state['analysis_data']:
        data = st.session_state['analysis_data']
        addr = st.session_state['address_cache']
        
        st.info(f"Viewing Analysis for: {addr}")
        
        if data.get('residual_land_value'):
            rlv = data['residual_land_value']
            st.success(f"**Maximum Walk-Away Price:** {rlv.get('max_walk_away_price', 'N/A')}  \n*{rlv.get('verdict', '')}*")

        max_units = int(data.get('schematic_data', {}).get('max_yield_units', 2))
        
        if max_units > 2:
            c1, c2, c3 = st.columns([2, 1, 1])
            c1.image(generate_wealth_chart(data.get('wealth_gap_data', {})), caption="Wealth Gap")
            c2.image(generate_schematic(data.get('schematic_data', {}), 2), caption="Duplex Option")
            c3.image(generate_schematic(data.get('schematic_data', {}), max_units), caption=f"Townhouses ({max_units})")
        else:
            c1, c2 = st.columns(2)
            c1.image(generate_wealth_chart(data.get('wealth_gap_data', {})), caption="Wealth Gap")
            c2.image(generate_schematic(data.get('schematic_data', {}), 2), caption="Duplex Option")
        
        st.divider()
        
        col_dl1, col_dl2 = st.columns(2)
        clean_name = addr.replace(",", "").replace(" ", "_")
        
        with col_dl1:
            doc_bytes = create_pro_report(data, addr)
            filename_doc = f"{clean_name}_{APP_VERSION}.docx"
            st.download_button(label=f"📄 Download Word Report", data=doc_bytes, file_name=filename_doc, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
        with col_dl2:
            excel_bytes = create_excel_report(data)
            filename_xls = f"{clean_name}_Financials.xlsx"
            st.download_button(label=f"📊 Download Live Excel Model", data=excel_bytes, file_name=filename_xls, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
